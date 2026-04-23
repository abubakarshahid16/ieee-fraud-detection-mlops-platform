from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Inches


ROOT = Path(".")
EVIDENCE_DIR = ROOT / "docs" / "evidence"
REPORT_MD = ROOT / "docs" / "research_report.md"
REQUIREMENTS_MD = ROOT / "docs" / "requirement_coverage.md"
OVERVIEW_MD = ROOT / "docs" / "project_overview.md"
DOCX_PATH = ROOT / "docs" / "research_report.docx"


def add_image_if_exists(document: Document, path: Path, width: float = 6.2) -> None:
    if path.exists():
        document.add_picture(str(path), width=Inches(width))


def main():
    document = Document()
    document.add_heading("Assignment 04 - Fraud Detection MLOps System", 0)
    document.add_paragraph(
        "This Word report was generated from the project artifacts using the real IEEE-CIS training data and includes step-by-step evidence screenshots."
    )

    document.add_heading("Overview", level=1)
    document.add_paragraph(REPORT_MD.read_text(encoding="utf-8")[:6000])

    document.add_heading("Complete Project Walkthrough", level=1)
    document.add_paragraph(OVERVIEW_MD.read_text(encoding="utf-8"))

    document.add_heading("Requirement Coverage", level=1)
    document.add_paragraph(REQUIREMENTS_MD.read_text(encoding="utf-8")[:8000])

    metrics_path = ROOT / "artifacts" / "metrics_summary.json"
    if metrics_path.exists():
        summary = json.loads(metrics_path.read_text(encoding="utf-8"))
        document.add_heading("Best Model Summary", level=1)
        document.add_paragraph(json.dumps(summary["best_metrics"], indent=2))
        document.add_paragraph(
            f"Best model: {summary['best_model']} | "
            f"Imbalance strategy: {summary['best_imbalance_strategy']} | "
            f"Cost strategy: {summary['best_cost_strategy']}"
        )

    document.add_heading("Step-by-Step Evidence", level=1)
    image_explanations = {
        "01_data_generation.png": "This screenshot summarizes the real IEEE-CIS dataset extraction and confirms the dataset size and fraud rate used in the final run.",
        "02_training_summary.png": "This screenshot shows the experiment summary, including the selected model, imbalance strategy, cost-sensitive configuration, and key evaluation metrics.",
        "03_confusion_matrix.png": "This screenshot visualizes the fraud-class confusion matrix so the recall and false-positive trade-off can be explained clearly.",
        "04_model_comparison.png": "This screenshot compares model families and imbalance/cost strategies side by side, highlighting the recall versus AUC trade-off.",
        "05_drift_summary.png": "This screenshot shows the strongest drifted features detected during the time-based drift analysis.",
        "06_kubeflow_compile.png": "This screenshot documents Kubeflow pipeline compilation and confirms the pipeline artifact generation step.",
        "07_k8s_validation.png": "This screenshot summarizes the live Kubernetes deployment evidence, including namespace, quota, service, and deployment rollout.",
        "08_tests_and_ci.png": "This screenshot summarizes build and packaging evidence, including Docker image creation, registry push, and runtime verification.",
        "09_api_prediction.png": "This screenshot captures live inference and metrics output from the running API service.",
        "10_monitoring_overview.png": "This screenshot summarizes the live Prometheus and Grafana stack, including provisioned dashboards and scraped metrics.",
        "11_alert_trigger.png": "This screenshot explains the live alert firing state and how it was turned into a retraining trigger artifact.",
        "12_submission_tree.png": "This screenshot provides the final submission structure included in the compressed package.",
        "13_transaction_distribution.png": "This screenshot shows a transaction amount distribution sample from the real data for quick dataset intuition.",
    }

    for name in [
        "01_data_generation.png",
        "02_training_summary.png",
        "03_confusion_matrix.png",
        "04_model_comparison.png",
        "05_drift_summary.png",
        "06_kubeflow_compile.png",
        "07_k8s_validation.png",
        "08_tests_and_ci.png",
        "09_api_prediction.png",
        "10_monitoring_overview.png",
        "11_alert_trigger.png",
        "12_submission_tree.png",
        "13_transaction_distribution.png",
    ]:
        image_path = EVIDENCE_DIR / name
        document.add_heading(name.replace(".png", "").replace("_", " "), level=2)
        document.add_paragraph(image_explanations.get(name, "Evidence screenshot."))
        add_image_if_exists(document, image_path)

    shap_path = ROOT / "artifacts" / "shap_summary.png"
    if shap_path.exists():
        document.add_heading("Explainability", level=1)
        document.add_paragraph(
            "This screenshot shows the SHAP-based explainability output used to justify why the fraud model predicts certain transactions as suspicious."
        )
        add_image_if_exists(document, shap_path)

    document.save(DOCX_PATH)
    print(str(DOCX_PATH.resolve()))


if __name__ == "__main__":
    main()
